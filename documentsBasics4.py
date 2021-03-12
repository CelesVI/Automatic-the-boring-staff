import docx
d = docx.Document(r'C:\Users\Bravin\Desktop\Python\Automatic the boring stuff\maravillas.docx')
print(d.paragraphs[1].text)
p = d.paragraphs[1].text
p.runs
